import hyperdiv as hd
import itertools
import sqlite3
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import webbrowser
from  datetime import date
from  datetime import datetime
from docx import Document
from docx.shared import Pt
import xlsxwriter
from  datetime import date
from  datetime import datetime
import csv


#mainrouter = hd.router()
foundIt = True
foundDialog = True
listkeys = []
listvalues = []
listdialog = []

def CreateSpreadsheet(tempvaluesPriority, tempvaluesPriorityDesc):
     
     d = datetime
     currDateTime = str(d.now())
     currDate = currDateTime[0:10]
     currTime = currDateTime[11:]
     currTime = currTime.replace(":","_")
     fileName = "Prior"+currDate+"_"+currTime+".xlsx"
     fullfilename = './Spreadsheet/'+fileName
     workbook = xlsxwriter.Workbook(fullfilename)
     worksheet = workbook.add_worksheet()
     bold = workbook.add_format({"bold": True})
     worksheet.set_column("B:B", 40)
     worksheet.set_column("C:C", 25)
     worksheet.set_column("D:D", 25)
     worksheet.set_column("E:E", 12)
     worksheet.set_column("F:F", 12)
     worksheet.write("A1", "Id", bold)
     worksheet.write("B1", "Description", bold)

     for i in range(0,len(tempvaluesPriority)):
          j = i + 1

          worksheet.write(str("A"+str(j+1)),str(tempvaluesPriority[i]))
          worksheet.write(str("B"+str(j+1)),str(tempvaluesPriorityDesc[i]))
    
     workbook.close()
     webopen = "http://localhost:8000/"+fullfilename
     webbrowser.open(webopen, new=0, autoraise=True)


def CreateCSV(tempvaluesPriority, tempvaluesPriorityDesc):
     d = datetime
     currDateTime = str(d.now())
     currDate = currDateTime[0:10]
     currTime = currDateTime[11:]
     currTime = currTime.replace(":","_")
     fileName = "Prior"+currDate+"_"+currTime+".csv"
     fullfilename = './csv/'+fileName
     header = ['Priority', 'Description']


     with open(fullfilename, "w", newline='') as f:
          writer = csv.writer(f)
          writer.writerow(header) 
          for i in range(0,len(tempvaluesPriority)):
               data = []
               data.append(str(tempvaluesPriority[i]))
               data.append(str(tempvaluesPriorityDesc[i]))           
               writer.writerow(data)                    
     f.close() 
     webopen = "http://localhost:8000/"+fullfilename
     print(webopen)
     webbrowser.open(webopen, new=0, autoraise=True)

def CreateWord(tempvaluesPriority, tempvaluesPriorityDesc):
     d = datetime
     currDateTime = str(d.now())
     currDate = currDateTime[0:10]
     currDate = currDate.replace("-","")
     currTime = currDateTime[11:]
     currTime = currTime.replace(":","")
     currTime = currTime.replace(".","")
     fileName = "Prior"+currDate+"_"+currTime+".docx"
     FullFileName = './word/'+fileName
     doc = Document()
     doc.add_heading('Priority', level=1)
# Add a table with 3 rows and 3 columns
     table = doc.add_table(rows=7, cols=7)
     table.style = 'LightShading-Accent1'
     table.width = doc.sections[0].page_width * 0.8
     #table.columns[1].style.font=8   
# Populate the table
     table.cell(0, 0).text = 'Priority'
     table.cell(0, 1).text = 'Description'
    
     for i in range(0,len(tempvaluesPriority)):
          j = i + 1
          table.cell(j,0).text = str(tempvaluesPriority[i])
          table.cell(j,1).text = str(tempvaluesPriorityDesc[i])  
     for row in table.rows:
       for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(8)
        #c.showPage()
        #tc = table.cell(0, 0).paragraphs[0].runs
        #tc[0].font.size = Pt(8)

# Save the document

     webopen = "http://localhost:8000/"+FullFileName

     doc.save(FullFileName)   
     webbrowser.open(webopen, new=0, autoraise=True)  

def CreatePdf(tempvaluesPriority, tempvaluesPriorityDesc):
     d = datetime
     currDateTime = str(d.now())
     currDate = currDateTime[0:10]
     currDate = currDate.replace("-","")
     currTime = currDateTime[11:]
     currTime = currTime.replace(":","")
     currTime = currTime.replace(".","")
     fileName = "Prior"+currDate+"_"+currTime+".pdf"
     FullFileName = './pdf/'+fileName
     webopen = "http://localhost:8000/"+FullFileName
     c = canvas.Canvas(FullFileName, pagesize=letter)
     w, h = letter
     height = int(h)
     c.setFont("Helvetica", 14)
     c.drawString(200, h - 40, "Priority Reporting")
     c.setFont("Helvetica", 12)
     c.drawString(30, height- 70, "Id")
     c.drawString(60, height - 70, "Description")
 

     for i in range(0, len(tempvaluesPriority)):
          heightDetail= (height-70)-(20+(i*20))
          c.drawString(30,heightDetail,str(tempvaluesPriority[i]))
          c.drawString(60,heightDetail,str(tempvaluesPriorityDesc[i]))

     c.save()
     webbrowser.open(webopen, new=0, autoraise=True)


  
def deletePriority(id):
     con = sqlite3.connect("requirements.db") 
     cur = con.cursor()
     try:
        sql = "DELETE FROM reqPriority where Priority = ?"
        cur.execute(sql, (id,))  
     except sqlite3.OperationalError as e:
          print(e) 
     finally:       
        con.commit()
        con.close()

def getPriorities():
     con = sqlite3.connect("requirements.db") 
     cur = con.cursor()  
     cur.execute("SELECT Priority,PriorityDesc from reqPriority")  
     userall =  cur.fetchall()
     con.commit()
     con.close()
     return userall



def insertPriority(Priority, PriorityDesc):
     con = sqlite3.connect("requirements.db") 
     cur = con.cursor()
     try:
        sql = "INSERT INTO reqPriority VALUES( ?,?)"
        cur.execute(sql, (Priority, PriorityDesc ))  
     except sqlite3.OperationalError as e:
          print(e) 
     finally:       
        con.commit()
        con.close()

def UpdatePriority(Priority, PriorityDesc):
     print("Desc"+PriorityDesc)
     print("Pror"+Priority)
     con = sqlite3.connect("requirements.db") 
     cur = con.cursor()
     try:
        sql = "Update reqPriority  set PriorityDesc = ? WHERE Priority = ?"
        cur.execute(sql, (PriorityDesc, Priority,))  
     except sqlite3.OperationalError as e:
          print(e) 
     finally:       
        con.commit()
        con.close()

def Editaction(idName):
     Editdialog = hd.dialog("Edit "+ Title)

     merged2 = []


     with Editdialog:
              ErrorArray = []
              ErrorArray.clear()
              merged = listvalues
              index = 0
              index = listvalues[0].index(idName)   

              loop = 0
             
              with hd.form( ) as form:    
                        hd.text(listvalues[0][0])
                        for i in range(1,len(listkeys)):                           
                             with hd.scope(i):                
                                form.text_input(listkeys[i], value=listvalues[i][index] )
                        with hd.hbox(gap=1):
                            form.submit_button(variant="primary")
                            form.reset_button()
                        if form.submitted:  
                            if len(ErrorArray) == 0:   
                                #print(listvalues[0][index]) 
                                for i in range(1,len(listkeys)):
                                    with hd.scope(i):               
                                         listvalues[i][index]  = form.form_data[listkeys[i]]                      
                                         datatabledict.update({listkeys[i]:listvalues[i]})   
                                #print(listvalues[0][index])         
                                UpdatePriority(listvalues[0][index], listvalues[1][index] )         
                                Editdialog.opened = False
                            else:
                                Editdialog.opened = True 
              

     with hd.hbox(padding=0.3, gap=0.3):
        with hd.tooltip(f"Edit {idName}"):
            edit=hd.button(
                prefix_icon="pencil",
                size="small",
                outline=True
            )
        with hd.tooltip(f"Delete {idName}"):
            trash= hd.button(
                prefix_icon="trash",
                size="small",
                outline=True
            )
        if edit.clicked:
            Editdialog.opened = True  
        if trash.clicked:
               index = listvalues[0].index(idName)
               try: 
                 for i in range(len(listkeys)):
                             with hd.scope(i):
                                    del listvalues[i][index]
                                    datatabledict.update({listkeys[i]:listvalues[i] })    
               except:
                   None 

               try:
                  deletePriority(idName)  
                  datatable.next_page()

               except:
                   None  

def TableValidation(value):

     if value == "":
          ErrorArray.append("Error") 
          return "please enter a value"
     

     




#UserDict, 3,True,"User","primary","bottom",types
def addTable(dict, perpage,action,Subject,addtype,placement,types):
    global dataTypes
    global datatabledict
    global Title
    Title = Subject
    hd.h3(Title)
    datatabledict = dict
    listvalues.clear()
    listkeys.clear()
    tempvaluesPriority = []
    tempvaluesPriorityDesc = []

    with hd.button_group():
        pdf=hd.button(
                "Pdf",
                variant="primary",
                prefix_icon="file-pdf",
                size="medium",
                outline=True
            )
        word=hd.button(
                "Word",
                variant="primary",
                prefix_icon="filetype-docx",
                size="medium",
                outline=True
            )
        spreadsheet=hd.button(
                "Spreadsheet",
                variant="primary",
                prefix_icon="file-spreadsheet",
                size="medium",
                outline=True
            )
        csv=hd.button(
                "Text File",
                variant="primary",
                prefix_icon="file-earmark-text",
                size="medium",
                outline=True
            )
        
    x= getPriorities()
    for rows in x: 
         tempvaluesPriority.append(rows [0]) 
         datatabledict.update({"Priority":tempvaluesPriority}) 
         tempvaluesPriorityDesc.append(rows [1]) 
         datatabledict.update({"PriorityDesc":tempvaluesPriorityDesc}) 


        
    
    for key in datatabledict.keys():
        listkeys.append(key)
    for value in datatabledict.values():    
        listvalues.append(value)
    idName = listkeys[0]
    

    global ErrorArray

    global datatable
    dataTypes = types    
    Adddialog = hd.dialog("Add "+ Title) 
    with Adddialog:
              global orgs
              orgs = []          
              ErrorArray = []  
              ErrorArray.clear()
              with hd.form() as form:          
                    for i in range(0,len(listkeys)):
                             with hd.scope(i):
                                try:
                                    hd.text(listkeys[i],background_color="gray-100") 
                                    form.text_input(name=listkeys[i],minlength=1,input_type=dataTypes[i], validation=TableValidation )                

                                except:
                                     None    

                    with hd.hbox(gap=1):
                            form.submit_button(variant="primary")
                            form.reset_button()
                    if form.submitted:
                      if len(ErrorArray) == 0:
                            datatabledict.update({listkeys[0]:listvalues[0]}) 
                            for i in range(1, len(listkeys)):
                                with hd.scope(i):
                                    listvalues[i].append(form.form_data[listkeys[i]])
                                    datatabledict.update({listkeys[i]:listvalues[i]}) 
                            insertPriority(form.form_data[listkeys[0]], form.form_data[listkeys[1]])
                            Adddialog.opened = False
                      else:
                         Adddialog.opened = True    

                             
  
     
    Editdialog = hd.dialog("Edit "+ Subject)   
    if placement == "top":
      with hd.button_group():
            with hd.button("add", variant=addtype, outline=True) as addbutton:
                hd.icon("plus", slot=addbutton.prefix)
            with hd.button("cancel", variant=addtype, outline=True) as cancelbutton:
                hd.icon("fullscreen-exit", slot=cancelbutton.prefix)        
    if action == True:
        try:
            datatable = hd.data_table(datatabledict, rows_per_page =perpage,id_column_name=idName, row_actions = Editaction)
             
        except:
             None    
    else:
        try:
            datatable = hd.data_table(datatabledict, rows_per_page =perpage)
        except:
             None    
    if placement == "bottom":
        with hd.button_group():
            with hd.button("add", variant=addtype, outline=True) as addbutton:
                hd.icon("plus", slot=addbutton.prefix)
            with hd.button("cancel", variant=addtype, outline=True) as cancelbutton:
                hd.icon("fullscreen-exit", slot=cancelbutton.prefix)        
        if addbutton.clicked:
            Adddialog.opened = True 
    if pdf.clicked:
         CreatePdf(tempvaluesPriority, tempvaluesPriorityDesc)       
    if word.clicked:
         CreateWord(tempvaluesPriority, tempvaluesPriorityDesc)       
    if spreadsheet.clicked:
         CreateSpreadsheet(tempvaluesPriority, tempvaluesPriorityDesc)                
    if csv.clicked:
         CreateCSV(tempvaluesPriority, tempvaluesPriorityDesc)                
    return dict 