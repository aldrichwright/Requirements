import hyperdiv as hd
import itertools
import sqlite3
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import requests
import webbrowser
from docx import Document
from docx.shared import Pt
import xlsxwriter
from  datetime import date
from  datetime import datetime
import csv

#mainrouter = hd.router()
foundIt = True
foundDialog = True
listkeys = []
listvalues = []
listdialog = []


  
def deleteRequirement(id):
     con = sqlite3.connect("requirements.db") 
     cur = con.cursor()
     try:
        sql = "DELETE FROM Requirements where ReqId = ?"
        cur.execute(sql, (id,))  
     except sqlite3.OperationalError as e:
          print(e) 
     finally:       
        con.commit()
        con.close()

def getRequirements():
     con = sqlite3.connect("requirements.db") 
     cur = con.cursor()  
     cur.execute("SELECT ReqId, ReqDesc, ReqCatType, ReqPrior, ReqInit, ReqRiskType, ReqDate, ReqBy from Requirements")  
     userall =  cur.fetchall()
     con.commit()
     con.close()
     return userall



def insertRequirements(ReqId,ReqDesc , ReqCatType, ReqPrior, ReqInit, ReqRiskType, ReqDate, ReqBy):
     con = sqlite3.connect("requirements.db") 
     cur = con.cursor()
     try:
        sql = "INSERT INTO Requirements VALUES( ?,?,?,?,?,?,?,?)"
        cur.execute(sql, (ReqId, ReqDesc, ReqCatType, ReqPrior, ReqInit, ReqRiskType, ReqDate, ReqBy ))  
     except sqlite3.OperationalError as e:
          print(e) 
     finally:       
        con.commit()
        con.close()

def UpdateRequirements(ReqDesc, ReqCatType, ReqPrior, ReqInit, ReqRiskType, ReqDate, ReqBy, ReqId,):       
     con = sqlite3.connect("requirements.db") 
     cur = con.cursor()
     try:
        sql = "Update Requirements set ReqDesc = ?, ReqCatType=?, ReqPrior=?, ReqInit=?, ReqRiskType=?,ReqDate=?, ReqBy=? WHERE ReqId = ?"
        cur.execute(sql, (ReqDesc, ReqCatType, ReqPrior, ReqInit, ReqRiskType, ReqDate, ReqBy, ReqId,))  
     except sqlite3.OperationalError as e:
          print(e) 
     finally:       
        con.commit()
        con.close()

def Editaction(id):
     Editdialog = hd.dialog("Edit "+ Title)

     merged2 = []


     with Editdialog:
              ErrorArray = []
              ErrorArray.clear()
              merged = listvalues
              index = 0
              index = listvalues[0].index(id)   

              loop = 0
             
              with hd.form( ) as form:    
                        for i in range(len(listkeys)):
                             with hd.scope(i):                
                                form.text_input(listkeys[i], value=listvalues[i][index] )
                        with hd.hbox(gap=1):
                            form.submit_button(variant="primary")
                            form.reset_button()
                        if form.submitted:  
                            if len(ErrorArray) == 0:  

                                for i in range(len(listkeys)):
                                    with hd.scope(i):               
                                         listvalues[i][index]  = form.form_data[listkeys[i]]                      
                                         datatabledict.update({listkeys[i]:listvalues[i]})        
                                UpdateRequirements(listvalues[1][index],listvalues[2][index],listvalues[3][index],listvalues[4][index],listvalues[5][index],listvalues[6][index],listvalues[7][index],listvalues[0][index] )         
                                Editdialog.opened = False
                            else:
                                Editdialog.opened = True 
              

     with hd.hbox(padding=0.3, gap=0.3):
        with hd.tooltip(f"Edit {id}"):
            edit=hd.button(
                prefix_icon="pencil",
                size="small",
                outline=True
            )
        with hd.tooltip(f"Delete {id}"):
            trash= hd.button(
                prefix_icon="trash",
                size="small",
                outline=True
            )
        if edit.clicked:
            Editdialog.opened = True  
        if trash.clicked:
               index = listvalues[0].index(id)
               try: 
                 for i in range(len(listkeys)):
                             with hd.scope(i):
                                    del listvalues[i][index]
                                    datatabledict.update({listkeys[i]:listvalues[i] })    
               except:
                   None 

               try:
                  deleteRequirement(id)  
                  datatable.next_page()

               except:
                   None  

def TableValidation(value):

     if value == "":
          ErrorArray.append("Error") 
          return "please enter a value"
     

def GetMaxRequirement():
     max_id = 0
     con = sqlite3.connect("requirements.db") 
     cur = con.cursor()  
     cur.execute("SELECT max(ReqId) FROM Requirements")  
     max_id_list =  cur.fetchone()
     max_id = max_id_list[0]   
     #print(max_id)
     con.commit()
     con.close()
     if max_id is None:
                max_id = 1
     else:
                max_id = max_id + 1       
     return max_id

def CreateSpreadsheet(tempvaluesReq, tempvaluesReqDesc,tempvaluesCat, tempvaluesInit, tempvaluesRisk, tempvaluesDate, tempvaluesBy):
     
     d = datetime
     currDateTime = str(d.now())
     currDate = currDateTime[0:10]
     currTime = currDateTime[11:]
     currTime = currTime.replace(":","_")
     fileName = "req"+currDate+"_"+currTime+".xlsx"
     fullfilename = './Spreadsheet/'+fileName
     workbook = xlsxwriter.Workbook(fullfilename)
     worksheet = workbook.add_worksheet()
     bold = workbook.add_format({"bold": True})
     worksheet.set_column("B:B", 40)
     worksheet.set_column("C:C", 25)
     worksheet.set_column("D:D", 25)
     worksheet.set_column("E:E", 12)
     worksheet.set_column("F:F", 12)
     worksheet.write("A1", "Id", bold)
     worksheet.write("B1", "Description", bold)
     worksheet.write("C1", "Category", bold)
     worksheet.write("D1", "Initiative", bold)
     worksheet.write("E1", "Risk", bold)
     worksheet.write("F1", "Date", bold)
     worksheet.write("G1", "By",bold)
     for i in range(0,len(tempvaluesReq)):
          j = i + 1

          worksheet.write(str("A"+str(j+1)),str(tempvaluesReq[i]))
          worksheet.write(str("B"+str(j+1)),str(tempvaluesReqDesc[i]))
          worksheet.write(str("C"+str(j+1)),str(tempvaluesCat[i]))
          worksheet.write(str("D"+str(j+1)),str(tempvaluesInit[i]))
          worksheet.write(str("E"+str(j+1)),str(tempvaluesRisk[i]))
          worksheet.write(str("F"+str(j+1)),str(tempvaluesDate[i]))
          worksheet.write(str("G"+str(j+1)),str(tempvaluesBy[i]))

     workbook.close()
     webopen = "http://localhost:8000/"+fullfilename
     webbrowser.open(webopen, new=0, autoraise=True)


def CreateCSV(tempvaluesReq,tempvaluesReqDesc, tempvaluesCat, tempvaluesInit, tempvaluesRisk, tempvaluesDate, tempvaluesBy):
     d = datetime
     currDateTime = str(d.now())
     currDate = currDateTime[0:10]
     currTime = currDateTime[11:]
     currTime = currTime.replace(":","_")
     fileName = "req"+currDate+"_"+currTime+".csv"
     fullfilename = './csv/'+fileName
     print(fullfilename)
     header = ['Id', 'Description', 'Category', 'Initiative', 'Risk','Date','By']


     with open(fullfilename, "w") as f:
          writer = csv.writer(f)
          writer.writerow(header) 
          for i in range(0,len(tempvaluesReq)):
               data = []
               data.append(str(tempvaluesReq[i]))
               data.append(str(tempvaluesReqDesc[i]))
               data.append(str(tempvaluesCat[i]))            
               data.append(str(tempvaluesInit[i]))            
               data.append(str(tempvaluesRisk[i]))            
               data.append(str(tempvaluesDate[i]))            
               data.append(str(tempvaluesBy[i]))            
               writer.writerow(data)                    
     f.close() 
     webopen = "http://localhost:8000/"+fullfilename
     print(webopen)
     webbrowser.open(webopen, new=0, autoraise=True)

def CreateWord(tempvaluesReq, tempvaluesReqDesc,tempvaluesCat, tempvaluesInit, tempvaluesRisk, tempvaluesDate, tempvaluesBy):
     d = datetime
     currDateTime = str(d.now())
     currDate = currDateTime[0:10]
     currDate = currDate.replace("-","")
     currTime = currDateTime[11:]
     currTime = currTime.replace(":","")
     currTime = currTime.replace(".","")
     fileName = "req"+currDate+"_"+currTime+".docx"
     FullFileName = './word/'+fileName
     doc = Document()
     doc.add_heading('Requirements', level=1)
# Add a table with 3 rows and 3 columns
     table = doc.add_table(rows=7, cols=7)
     table.style = 'LightShading-Accent1'
     table.width = doc.sections[0].page_width * 0.8
     #table.columns[1].style.font=8   
# Populate the table
     table.cell(0, 0).text = 'Id'
     table.cell(0, 1).text = 'Description'
     table.cell(0, 2).text = 'Category '
     table.cell(0, 3).text = 'Initiative'
     table.cell(0, 4).text = 'Risk'
     table.cell(0, 5).text = 'Date'
     table.cell(0, 6).text = 'By'
     for i in range(0,len(tempvaluesReq)):
          j = i + 1
          table.cell(j,0).text = str(tempvaluesReq[i])
          table.cell(j,1).text = str(tempvaluesReqDesc[i])  
          table.cell(j,2).text = str(tempvaluesCat[i])  
          table.cell(j,3).text = str(tempvaluesInit[i])  
          table.cell(j,4).text = str(tempvaluesRisk[i])  
          table.cell(j,5).text = str(tempvaluesDate[i])  
          table.cell(j,6).text = str(tempvaluesBy[i])  
     for row in table.rows:
       for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(8)
        #c.showPage()
        #tc = table.cell(0, 0).paragraphs[0].runs
        #tc[0].font.size = Pt(8)

# Save the document
     webopen = "http://localhost:8000/"+FullFileName
     doc.save(FullFileName)   
     webbrowser.open(webopen, new=0, autoraise=True)  

def CreatePdf(tempvaluesReq, tempvaluesReqDesc,tempvaluesCat, tempvaluesInit, tempvaluesRisk, tempvaluesDate, tempvaluesBy):
     d = datetime
     currDateTime = str(d.now())
     currDate = currDateTime[0:10]
     currDate = currDate.replace("-","")
     currTime = currDateTime[11:]
     currTime = currTime.replace(":","")
     currTime = currTime.replace(".","")
     fileName = "req"+currDate+"_"+currTime+".pdf"
     FullFileName = './pdf/'+fileName
     webopen = "http://localhost:8000/"+FullFileName
     c = canvas.Canvas(FullFileName, pagesize=letter)
     w, h = letter
     height = int(h)
     c.setFont("Helvetica", 14)
     c.drawString(200, h - 40, "Requirements Reporting")
     c.setFont("Helvetica", 12)
     c.drawString(30, height- 70, "Id")
     c.drawString(60, height - 70, "Description(30)")
     c.drawString(240, height - 70, "Category")
     c.drawString(300, height - 70, "Initiative")
     c.drawString(360, height - 70, "Risk")
     c.drawString(420, height - 70, "Date")    
     c.drawString(500, height - 70, "By")    

     for i in range(0, len(tempvaluesReq)):
          heightDetail= (height-70)-(20+(i*20))
          c.drawString(30,heightDetail,str(tempvaluesReq[i]))
          c.drawString(60,heightDetail,str(tempvaluesReqDesc[i])[:30])
          c.drawString(240,heightDetail,tempvaluesCat[i])
          c.drawString(300,heightDetail,tempvaluesInit[i])
          c.drawString(360,heightDetail,tempvaluesRisk[i])
          c.drawString(420,heightDetail,tempvaluesDate[i])
          c.drawString(500,heightDetail,tempvaluesBy[i])

     c.save()
     webbrowser.open(webopen, new=0, autoraise=True)



#UserDict, 3,True,"User","primary","bottom",types
def addTable(dict, perpage,action,Subject,addtype,placement,types):
    global dataTypes
    global datatabledict
    global Title
    Title = Subject
    hd.h3(Title)
    datatabledict = dict
    maxid =  GetMaxRequirement()
    listvalues.clear()
    listkeys.clear()
    tempvaluesReq = []
    tempvaluesReqDesc = []
    tempvaluesCat = []
    tempvaluesPrior = []
    tempvaluesInit = []
    tempvaluesRisk = []
    tempvaluesDate = []
    tempvaluesBy = []
    with hd.button_group():
        pdf=hd.button(
                "Pdf",
                variant="primary",
                prefix_icon="file-pdf",
                size="medium",
                outline=True
            )
        word= hd.button(
                "Word",
                variant="primary",
                prefix_icon="filetype-docx",
                size="medium",
                outline=True
            ) 
        spread= hd.button(
                "Spreadsheet",
                variant="primary",
                prefix_icon="file-spreadsheet",
                size="medium",
                outline=True
            )
        csv=hd.button(
                "Text File",
                variant="primary",
                prefix_icon="file-earmark-text",
                size="medium",
                outline=True
            ) 
    x= getRequirements()

    for rows in x:
         tempvaluesReq.append(rows [0]) 
         datatabledict.update({"ReqId":tempvaluesReq}) 
         tempvaluesReqDesc.append(rows [1]) 
         datatabledict.update({"Req Description":tempvaluesReqDesc}) 
         tempvaluesCat.append(rows [2]) 
         datatabledict.update({"Req Cat Type":tempvaluesCat}) 
         tempvaluesPrior.append(rows [3]) 
         datatabledict.update({"Req Prior":tempvaluesPrior}) 
         tempvaluesInit.append(rows [4]) 
         datatabledict.update({"Req Initiative":tempvaluesInit})         
         tempvaluesRisk.append(rows [5]) 
         datatabledict.update({"Req Risk":tempvaluesRisk})   
         tempvaluesDate.append(rows [6]) 
         datatabledict.update({"Req Date":tempvaluesDate})   
         tempvaluesBy.append(rows [7]) 
         datatabledict.update({"Req By":tempvaluesBy})   
    for key in datatabledict.keys():
        listkeys.append(key)
    for value in datatabledict.values():    
        listvalues.append(value)
    idName = listkeys[0]
    

    global ErrorArray

    global datatable
    dataTypes = types    
    Adddialog = hd.dialog("Add "+ Title) 
    with Adddialog:
              global orgs
              orgs = []          
              ErrorArray = []  
              ErrorArray.clear()
              with hd.form() as form: 
                    hd.text(listkeys[0],background_color="gray-100") 
                    hd.text(maxid)               
                    for i in range(1,len(listkeys)):
                             with hd.scope(i):
                                try:
                                    hd.text(listkeys[i],background_color="gray-100") 
                                    form.text_input(name=listkeys[i],minlength=1,input_type=dataTypes[i], validation=TableValidation )                

                                except:
                                     None    

                    with hd.hbox(gap=1):
                            form.submit_button(variant="primary")
                            form.reset_button()
                    if form.submitted:
                      if len(ErrorArray) == 0:
                            listvalues[0].append(maxid)
                            datatabledict.update({listkeys[0]:listvalues[0]}) 
                            for i in range(1, len(listkeys)):
                                with hd.scope(i):
                                    listvalues[i].append(form.form_data[listkeys[i]])
                                    datatabledict.update({listkeys[i]:listvalues[i]}) 
                            insertRequirements(maxid,form.form_data[listkeys[1]], form.form_data[listkeys[2]],form.form_data[listkeys[3]],form.form_data[listkeys[4]],form.form_data[listkeys[5]],form.form_data[listkeys[6]],form.form_data[listkeys[7]])
                            Adddialog.opened = False
                      else:
                         Adddialog.opened = True    

                             
  
     
    Editdialog = hd.dialog("Edit "+ Subject)   
    if placement == "top":
      with hd.button_group():
            with hd.button("add", variant=addtype, outline=True) as addbutton:
                hd.icon("plus", slot=addbutton.prefix)
            with hd.button("cancel", variant=addtype, outline=True) as cancelbutton:
                hd.icon("fullscreen-exit", slot=cancelbutton.prefix)        
    if action == True:
        try:
            datatable = hd.data_table(datatabledict, rows_per_page =perpage,id_column_name=idName, row_actions = Editaction)
             
        except:
             None    
    else:
        try:
            datatable = hd.data_table(datatabledict, rows_per_page =perpage)
        except:
             None    
    if placement == "bottom":
        with hd.button_group():
            with hd.button("add", variant=addtype, outline=True) as addbutton:
                hd.icon("plus", slot=addbutton.prefix)
            with hd.button("cancel", variant=addtype, outline=True) as cancelbutton:
                hd.icon("fullscreen-exit", slot=cancelbutton.prefix)  
        if cancelbutton.clicked:
             hd.location().go(path="/")          
        if addbutton.clicked:
            Adddialog.opened = True  
    if pdf.clicked:
         CreatePdf(tempvaluesReq,tempvaluesReqDesc, tempvaluesCat, tempvaluesInit, tempvaluesRisk, tempvaluesDate, tempvaluesBy)       
    if word.clicked:
         CreateWord(tempvaluesReq,tempvaluesReqDesc, tempvaluesCat, tempvaluesInit, tempvaluesRisk, tempvaluesDate, tempvaluesBy)           
    if spread.clicked:
         CreateSpreadsheet(tempvaluesReq,tempvaluesReqDesc, tempvaluesCat, tempvaluesInit, tempvaluesRisk, tempvaluesDate, tempvaluesBy)                
    if csv.clicked:
         CreateCSV(tempvaluesReq,tempvaluesReqDesc, tempvaluesCat, tempvaluesInit, tempvaluesRisk, tempvaluesDate, tempvaluesBy)                     
    return dict 